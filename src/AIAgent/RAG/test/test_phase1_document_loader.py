from datetime import datetime
from pathlib import Path
from types import SimpleNamespace

import pytest

from phase1_document_loader import DocumentLoader


def text_block(text, bbox=(0, 0, 100, 20)):
    return {
        "type": 0,
        "bbox": bbox,
        "lines": [{"spans": [{"text": text}]}],
    }


class FakePage:
    def __init__(self, blocks):
        self.blocks = blocks

    def get_text(self, fmt, sort=True):
        assert fmt == "dict"
        return {"blocks": self.blocks}

    def find_tables(self):
        return SimpleNamespace(tables=[])


class FakePdf:
    is_encrypted = False

    def __init__(self, pages):
        self.pages = pages
        self.page_count = len(pages)

    def __getitem__(self, index):
        return self.pages[index]

    def authenticate(self, password):
        return True

    def close(self):
        pass


def test_document_metadata_has_stable_base_fields(tmp_path):
    filepath = tmp_path / "Sample.TXT"
    filepath.write_text("hello metadata", encoding="utf-8")

    loaded_documents = DocumentLoader().load_document(filepath)
    loaded = loaded_documents[0]

    assert len(loaded_documents) == 1
    assert loaded.metadata["source"] == "Sample.TXT"
    assert loaded.metadata["filename"] == "Sample.TXT"
    assert loaded.metadata["filepath"] == str(filepath.resolve())
    assert loaded.metadata["extension"] == ".txt"
    assert loaded.metadata["mime_type"] == "text/plain"
    assert loaded.metadata["loader"] == "text"
    assert loaded.metadata["char_count"] == len("hello metadata")
    assert loaded.metadata["file_size"] == filepath.stat().st_size
    assert loaded.metadata["part_index"] == 0
    assert loaded.metadata["part_total"] == 1
    assert loaded.metadata["part_type"] == "text"
    assert isinstance(loaded.metadata["document_id"], str)
    assert len(loaded.metadata["document_id"]) == 16
    assert datetime.fromisoformat(str(loaded.metadata["file_modified_at"]))


def test_document_repr():
    from phase1_document_loader import Document

    # 1. Test standard document
    doc1 = Document(content="hello world", metadata={"source": "test.txt"})
    assert repr(doc1) == "Document(source=test.txt, len=11, preview='hello world...')"

    # 2. Test multi-part / page document
    doc2 = Document(
        content="page content",
        metadata={
            "source": "test.pdf",
            "page": 3,
            "part_index": 2,
            "part_total": 5,
        }
    )
    assert repr(doc2) == "Document(source=test.pdf, page=3, part=3/5, len=12, preview='page content...')"


def test_docx_loader_extracts_paragraphs_and_tables(tmp_path):
    from docx import Document as DocxDocument

    filepath = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("第一段内容")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "含义"
    table.cell(1, 0).text = "RAG"
    table.cell(1, 1).text = "检索增强生成"
    doc.add_paragraph("最后一段")
    doc.save(filepath)

    loaded = DocumentLoader().load_document(filepath)

    assert [doc.metadata["part_type"] for doc in loaded] == [
        "paragraph",
        "table",
        "paragraph",
    ]
    assert loaded[0].content == "第一段内容"
    assert "| 字段 | 含义 |" in loaded[1].content
    assert "| RAG | 检索增强生成 |" in loaded[1].content
    assert loaded[2].content == "最后一段"
    assert [doc.metadata["part_index"] for doc in loaded] == [0, 1, 2]
    assert all(doc.metadata["part_total"] == 3 for doc in loaded)


def test_csv_loader_formats_rows_and_handles_empty_cells(tmp_path):
    filepath = tmp_path / "sample.csv"
    filepath.write_text(
        "name,age,note\n"
        "Alice,30,\n"
        "Bob,,hello|pipe\n"
        "Short\n"
        "Long,1,2,3\n",
        encoding="utf-8",
    )

    loaded = DocumentLoader().load_document(filepath)[0]

    assert "| name | age | note |  |" in loaded.content
    assert "| Alice | 30 |  |  |" in loaded.content
    assert "| Bob |  | hello\\|pipe |  |" in loaded.content
    assert "| Short |  |  |  |" in loaded.content
    assert "| Long | 1 | 2 | 3 |" in loaded.content


def test_csv_loader_sniffs_delimiter_and_preserves_quoted_newlines(tmp_path):
    filepath = tmp_path / "sample_semicolon.csv"
    filepath.write_text('name;note\nAlice;"line 1\nline 2"\n', encoding="utf-8")

    loaded = DocumentLoader().load_document(filepath)[0]

    assert "| name | note |" in loaded.content
    assert "| Alice | line 1<br>line 2 |" in loaded.content


def test_json_loader_formats_json_and_preserves_chinese(tmp_path):
    filepath = tmp_path / "sample.json"
    filepath.write_text('{"name":"张三","items":[1,2]}', encoding="utf-8")

    loaded = DocumentLoader().load_document(filepath)[0]

    assert '"name": "张三"' in loaded.content
    assert '"items": [' in loaded.content
    assert "\\u5f20" not in loaded.content


def test_invalid_json_raises_clear_error(tmp_path):
    filepath = tmp_path / "broken.json"
    filepath.write_text("{bad json", encoding="utf-8")

    with pytest.raises(ValueError, match="无法解析 JSON 文件 broken.json"):
        DocumentLoader().load_document(filepath)


def test_empty_csv_is_rejected_by_load_document(tmp_path):
    filepath = tmp_path / "empty.csv"
    filepath.write_text("", encoding="utf-8")

    with pytest.raises(ValueError, match="内容为空"):
        DocumentLoader().load_document(filepath)


def test_pdf_loader_creates_one_document_per_page(tmp_path, monkeypatch):
    import fitz

    filepath = tmp_path / "sample.pdf"
    filepath.write_bytes(b"%PDF fake")
    fake_pdf = FakePdf(
        [
            FakePage([text_block("First page has enough native text.")]),
            FakePage([text_block("Second page has enough native text.")]),
        ]
    )
    monkeypatch.setattr(fitz, "open", lambda path: fake_pdf)

    loaded = DocumentLoader().load_document(filepath)

    assert len(loaded) == 2
    assert loaded[0].metadata["page"] == 1
    assert loaded[1].metadata["page"] == 2
    assert loaded[0].metadata["part_type"] == "page"
    assert loaded[1].metadata["part_type"] == "page"
    assert loaded[0].metadata["part_index"] == 0
    assert loaded[1].metadata["part_index"] == 1
    assert all(doc.metadata["part_total"] == 2 for doc in loaded)
    assert "First page" in loaded[0].content
    assert "Second page" in loaded[1].content
